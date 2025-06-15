from pathlib import Path
from typing import List, Tuple
import logging

logger = logging.getLogger(__name__)

def pdf_to_text(path: str) -> str:
    """Extract text from a PDF file."""
    try:
        from pdfminer.high_level import extract_text
        return extract_text(path)
    except Exception as e:
        logger.error(f"Failed to parse PDF {path}: {e}")
        return ""

def docx_to_text(path: Path) -> str:
    """Extract text from DOCX file"""
    try:
        from docx import Document
        doc = Document(str(path))
        text_parts = []
        
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text.strip())
        
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text_parts.append(cell.text.strip())
        
        return '\n'.join(text_parts)
    except Exception as e:
        logger.error(f"Failed to parse DOCX {path}: {e}")
        return ""

def load_text_file(path: Path) -> str:
    """Load text file with encoding detection"""
    try:
        return path.read_text(encoding="utf-8")
    except UnicodeDecodeError:
        try:
            return path.read_text(encoding="latin-1")
        except Exception as e:
            logger.error(f"Failed to read text file {path}: {e}")
            return ""

def load_resume(path: Path) -> Tuple[str, bool]:
    """
    Load resume from file
    Returns: (text_content, success)
    """
    try:
        if path.suffix.lower() == ".pdf":
            text = pdf_to_text(str(path))
        elif path.suffix.lower() in {".docx", ".doc"}:
            text = docx_to_text(path)
        else:
            text = load_text_file(path)
        
        return text.strip(), bool(text.strip())
    except Exception as e:
        logger.error(f"Failed to load resume {path}: {e}")
        return "", False

def load_resumes(paths: List[Path]) -> List[Tuple[str, str]]:
    """
    Load multiple resumes
    Returns: List of (filename, content) tuples for successful loads
    """
    results = []
    for path in paths:
        content, success = load_resume(path)
        if success:
            results.append((path.name, content))
        else:
            logger.warning(f"Failed to load resume: {path}")
    
    return results

def is_supported_file(path: Path) -> bool:
    """Check if file type is supported"""
    return path.suffix.lower() in {'.pdf', '.docx', '.doc', '.txt'}